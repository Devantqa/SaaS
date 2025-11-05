import os
import datetime
from docx import Document
from docx.shared import Inches
from PIL import Image  # noqa: F401  (Pillow required by python-docx for image handling)

class DocxReport:
    def __init__(self, out_path="artifacts/TestSummary.docx"):
        self.out_path = out_path
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
        self.doc = Document()
        self._add_meta()

    def _add_meta(self):
        self.doc.core_properties.author = "Automation"
        self.doc.core_properties.created = datetime.datetime.now()

    def add_title(self, text):
        self.doc.add_heading(text, 0)

    def add_step(self, title, description=None, screenshot_path=None):
        self.doc.add_heading(title, level=2)
        if description:
            self.doc.add_paragraph(description)
        if screenshot_path and os.path.exists(screenshot_path):
            self.doc.add_picture(screenshot_path, width=Inches(6.5))

    def add_info(self, text):
        self.doc.add_paragraph(text)

    def save(self):
        self.doc.save(self.out_path)

class ScreenshotHelper:
    def __init__(self, driver, report: DocxReport):
        self.driver = driver
        self.report = report
        self.folder = os.path.join("artifacts", "screenshots")
        os.makedirs(self.folder, exist_ok=True)

    def capture(self, name, add_to_report=True, description=None):
        safe = "".join([c if c.isalnum() or c in ("_","-") else "_" for c in name])
        path = os.path.join(
            self.folder,
            f"{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}_{safe}.png"
        )
        self.driver.save_screenshot(path)
        if add_to_report:
            self.report.add_step(title=name, description=description, screenshot_path=path)
        return path
